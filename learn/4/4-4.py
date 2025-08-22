import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from datetime import datetime, timedelta

# ----------------------
# 关键修复：设置Matplotlib中文字体（放在绘图前）
# ----------------------
plt.rcParams["font.family"] = ["SimHei", "WenQuanYi Micro Hei", "Heiti TC"]  # 支持中文
plt.rcParams["axes.unicode_minus"] = False  # 解决负号显示问题（避免方块）


# 1. 模拟数据（生产数据）
dates = [datetime.now() - timedelta(days=i) for i in range(7, 0, -1)]
data = {
    "日期": [d.strftime("%Y-%m-%d") for d in dates],
    "生产数量(根)": [120, 135, 128, 140, 132, 145, 150],
    "合格率(%)": [98.2, 97.8, 98.5, 99.0, 98.7, 99.2, 99.5]
}
df = pd.DataFrame(data)


# 2. 生成趋势图表（已支持中文）
plt.figure(figsize=(8, 4))

# 生产数量（柱状图）
plt.bar(df["日期"], df["生产数量(根)"], color="#4CAF50", alpha=0.7, label="生产数量")
plt.ylabel("生产数量(根)")  # 中文标签

# 合格率（折线图，共享x轴）
ax2 = plt.twinx()
ax2.plot(df["日期"], df["合格率(%)"], color="#2196F3", marker="o", label="合格率")
ax2.set_ylabel("合格率(%)")  # 中文标签

plt.title("核燃料棒周生产趋势")  # 中文标题
plt.tight_layout()  # 自动调整布局，避免标签重叠
plt.savefig("production_trend.png")  # 保存图表（此时中文已正常显示）


# 3. 生成Word报告
doc = Document()
doc.add_heading("核燃料生产周报告", 0)

# 添加基本信息
doc.add_paragraph(f"报告周期：{df['日期'].min()} 至 {df['日期'].max()}")
doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")

# 添加汇总数据
doc.add_heading("一、周汇总指标", level=1)
table = doc.add_table(rows=2, cols=2)
table.cell(0, 0).text = "总生产数量"
table.cell(0, 1).text = f"{df['生产数量(根)'].sum()} 根"
table.cell(1, 0).text = "平均合格率"
table.cell(1, 1).text = f"{df['合格率(%)'].mean().round(2)}%"

# 添加图表
doc.add_heading("二、生产趋势分析", level=1)
doc.add_paragraph("本周生产数量与合格率均呈上升趋势，合格率较上周提升0.5%。")
doc.add_picture("production_trend.png", width=Inches(6))  # 插入修复后的图表

# 保存报告
report_path = f"核燃料生产周报告_{df['日期'].max()}.docx"
doc.save(report_path)
print(f"报告生成成功：{report_path}")